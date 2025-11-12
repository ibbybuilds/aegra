"""File content extraction utilities for multimodal messages.

Handles extraction of text content from various file formats uploaded by users,
including PDFs, DOCX, spreadsheets, and code files.
"""

import base64
import io
import logging
from typing import Any

logger = logging.getLogger(__name__)

# Maximum characters to extract from files (to prevent slowdowns)
MAX_CONTENT_LENGTH = 50000  # ~50KB of text
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB max file size for processing


def extract_text_from_base64(
    data: str, mime_type: str, filename: str = "unknown"
) -> str:
    """Extract text content from base64-encoded files.

    Args:
        data: Base64-encoded file content
        mime_type: MIME type of the file
        filename: Name of the file (for context)

    Returns:
        Extracted text content or error message
    """
    try:
        # Quick size check before decoding
        estimated_size = len(data) * 3 // 4  # Base64 overhead
        if estimated_size > MAX_FILE_SIZE:
            return f"[File too large: {filename} (~{estimated_size // 1024 // 1024}MB). Please upload files < 5MB for text extraction]"

        decoded = base64.b64decode(data)

        # Text-based files (can be decoded directly)
        if mime_type.startswith("text/") or mime_type in [
            "application/json",
            "application/xml",
            "application/javascript",
            "application/typescript",
        ]:
            text = decoded.decode("utf-8", errors="replace")
            return text[:MAX_CONTENT_LENGTH] + (
                "\n\n[...content truncated for length...]"
                if len(text) > MAX_CONTENT_LENGTH
                else ""
            )

        # CSV files
        if mime_type == "text/csv":
            import csv

            text_io = io.StringIO(decoded.decode("utf-8", errors="replace"))
            reader = csv.reader(text_io)
            rows = list(reader)
            content = "\n".join([", ".join(row) for row in rows[:100]])  # Limit rows
            return content[:MAX_CONTENT_LENGTH] + (
                "\n\n[...rows truncated...]" if len(rows) > 100 else ""
            )

        # PDF files
        if mime_type == "application/pdf":
            try:
                from pypdf import PdfReader

                pdf_io = io.BytesIO(decoded)
                reader = PdfReader(pdf_io)
                text_parts = []
                # Limit to first 10 pages for speed
                pages_to_extract = min(10, len(reader.pages))
                for page_num, page in enumerate(reader.pages[:pages_to_extract]):
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- Page {page_num + 1} ---\n{text}")

                full_text = (
                    "\n\n".join(text_parts) if text_parts else "[No text found in PDF]"
                )
                truncated = full_text[:MAX_CONTENT_LENGTH]
                suffix = ""
                if len(full_text) > MAX_CONTENT_LENGTH:
                    suffix = "\n\n[...content truncated for length...]"
                elif len(reader.pages) > pages_to_extract:
                    suffix = f"\n\n[...{len(reader.pages) - pages_to_extract} more pages not shown...]"
                return truncated + suffix
            except ImportError:
                logger.warning("pypdf not installed, cannot extract PDF content")
                return f"[PDF file: {filename}. Install pypdf to extract text content]"
            except Exception as e:
                logger.error(f"Error extracting PDF content: {e}")
                return f"[Error extracting PDF content: {str(e)}]"

        # DOCX files
        if (
            mime_type
            == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ):
            try:
                from docx import Document

                docx_io = io.BytesIO(decoded)
                doc = Document(docx_io)
                # Limit to first 50 paragraphs
                paragraphs = [p.text for p in doc.paragraphs[:50]]
                full_text = "\n\n".join(paragraphs)
                truncated = full_text[:MAX_CONTENT_LENGTH]
                suffix = ""
                if len(full_text) > MAX_CONTENT_LENGTH:
                    suffix = "\n\n[...content truncated for length...]"
                elif len(doc.paragraphs) > 50:
                    suffix = f"\n\n[...{len(doc.paragraphs) - 50} more paragraphs not shown...]"
                return truncated + suffix
            except ImportError:
                logger.warning("python-docx not installed, cannot extract DOCX content")
                return f"[DOCX file: {filename}. Install python-docx to extract text content]"
            except Exception as e:
                logger.error(f"Error extracting DOCX content: {e}")
                return f"[Error extracting DOCX content: {str(e)}]"

        # XLSX files
        if (
            mime_type
            == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        ):
            try:
                import openpyxl

                xlsx_io = io.BytesIO(decoded)
                wb = openpyxl.load_workbook(xlsx_io, read_only=True, data_only=True)
                sheets_content = []
                for sheet_name in wb.sheetnames[:3]:  # Limit to 3 sheets
                    ws = wb[sheet_name]
                    rows = []
                    for row_num, row in enumerate(ws.iter_rows(values_only=True), 1):
                        if row_num > 50:  # Limit to 50 rows per sheet
                            rows.append(
                                f"[...{ws.max_row - 50} more rows not shown...]"
                            )
                            break
                        rows.append(
                            ", ".join(
                                [str(cell) if cell is not None else "" for cell in row]
                            )
                        )
                    sheets_content.append(
                        f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows)
                    )
                full_text = "\n\n".join(sheets_content)
                return full_text[:MAX_CONTENT_LENGTH] + (
                    "\n\n[...content truncated for length...]"
                    if len(full_text) > MAX_CONTENT_LENGTH
                    else ""
                )
            except ImportError:
                logger.warning("openpyxl not installed, cannot extract XLSX content")
                return f"[XLSX file: {filename}. Install openpyxl to extract content]"
            except Exception as e:
                logger.error(f"Error extracting XLSX content: {e}")
                return f"[Error extracting XLSX content: {str(e)}]"

        # Unknown file type
        return f"[File: {filename} ({mime_type}). Content extraction not supported for this file type]"

    except Exception as e:
        logger.error(f"Error decoding base64 file content: {e}", exc_info=True)
        return f"[Error reading file: {str(e)}]"


def process_multimodal_content(content: list[dict[str, Any]]) -> str:
    """Process multimodal message content and extract text from all blocks.

    Args:
        content: List of content blocks from a message

    Returns:
        Combined text representation of all content blocks
    """
    import time

    start_time = time.time()

    if not isinstance(content, list):
        return str(content)

    text_parts = []
    file_count = 0

    for block in content:
        if not isinstance(block, dict):
            continue

        block_type = block.get("type")

        # Text blocks
        if block_type == "text":
            text_parts.append(block.get("text", ""))

        # Image blocks (just note presence)
        elif block_type == "image":
            metadata = block.get("metadata", {})
            name = metadata.get("name", "unknown")
            text_parts.append(f"[Image: {name}]")

        # File blocks (extract content)
        elif block_type == "file":
            file_count += 1
            source_type = block.get("source_type")
            if source_type == "base64":
                data = block.get("data", "")
                mime_type = block.get("mime_type", "")
                metadata = block.get("metadata", {})
                filename = metadata.get("filename", "unknown")

                extracted = extract_text_from_base64(data, mime_type, filename)
                text_parts.append(f"\n--- File: {filename} ---\n{extracted}\n")

        # Text-plain blocks (text/code files)
        elif block_type == "text-plain":
            file_count += 1
            source_type = block.get("source_type")
            if source_type == "base64":
                data = block.get("data", "")
                mime_type = block.get("mime_type", "text/plain")
                metadata = block.get("metadata", {})
                filename = metadata.get("filename", "unknown")

                extracted = extract_text_from_base64(data, mime_type, filename)
                text_parts.append(f"\n--- File: {filename} ---\n{extracted}\n")

    elapsed = time.time() - start_time
    if file_count > 0:
        logger.info(f"Processed {file_count} file(s) in {elapsed:.2f}s")

    return "\n".join(text_parts)
